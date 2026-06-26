"""
Document parsers for each supported format.
Returns (raw_text, metadata_dict) for each input type.
"""
import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

import httpx
from html2text import HTML2Text

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    text: str
    title: str = ""
    author: str = ""
    page_count: int = 1
    language: str = "en"
    source_type: str = "txt"
    extra: dict = field(default_factory=dict)


# ── PDF ──────────────────────────────────────────────────────────────────────
def parse_pdf(content: bytes) -> ParsedDocument:
    import pdfplumber

    pages_text: list[str] = []
    title = ""
    author = ""

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        meta = pdf.metadata or {}
        title = meta.get("Title", "") or ""
        author = meta.get("Author", "") or ""
        page_count = len(pdf.pages)

        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text.strip())

    raw = "\n\n".join(pages_text)
    return ParsedDocument(
        text=raw,
        title=title,
        author=author,
        page_count=page_count,
        source_type="pdf",
    )


# ── DOCX ─────────────────────────────────────────────────────────────────────
def parse_docx(content: bytes) -> ParsedDocument:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Try to extract core properties
    try:
        props = doc.core_properties
        title = props.title or ""
        author = props.author or ""
    except Exception:
        title = author = ""

    return ParsedDocument(
        text="\n\n".join(paragraphs),
        title=title,
        author=author,
        page_count=1,
        source_type="docx",
    )


# ── TXT / MD ─────────────────────────────────────────────────────────────────
def parse_text(content: bytes, source_type: str = "txt") -> ParsedDocument:
    text = content.decode("utf-8", errors="replace")
    return ParsedDocument(text=text, source_type=source_type)


# ── URL ───────────────────────────────────────────────────────────────────────
def parse_url(url: str, depth: int = 0) -> ParsedDocument:
    """Fetch URL, convert HTML to markdown text."""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (compatible; AdaptiveRAG/1.0; +https://github.com/adaptive-rag)"
        )
    }

    with httpx.Client(timeout=30, follow_redirects=True) as client:
        resp = client.get(url, headers=headers)
        resp.raise_for_status()
        html = resp.text

    converter = HTML2Text()
    converter.ignore_links = False
    converter.ignore_images = True
    converter.body_width = 0  # no line wrapping
    text = converter.handle(html)

    # Extract domain as rough title
    parsed = urlparse(url)
    title = parsed.netloc

    return ParsedDocument(
        text=text,
        title=title,
        source_type="url",
        extra={"url": url, "depth": depth},
    )


# ── Language detection ────────────────────────────────────────────────────────
def detect_language(text: str) -> str:
    try:
        from langdetect import detect
        return detect(text[:2000])  # use first 2000 chars for speed
    except Exception:
        return "en"


# ── Cleaning ─────────────────────────────────────────────────────────────────
def clean_text(text: str) -> str:
    """
    Remove boilerplate, normalize whitespace, fix common encoding artefacts.
    """
    import re

    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove page-number-only lines like "- 5 -" or "5"
    text = re.sub(r"^\s*[-–—]?\s*\d+\s*[-–—]?\s*$", "", text, flags=re.MULTILINE)
    # Fix Windows line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Strip leading/trailing whitespace per line
    lines = [ln.rstrip() for ln in text.splitlines()]
    return "\n".join(lines).strip()


# ── Dispatcher ────────────────────────────────────────────────────────────────
def parse_document(
    content: bytes, filename: str, source_type: Optional[str] = None
) -> ParsedDocument:
    ext = Path(filename).suffix.lower().lstrip(".")
    stype = source_type or ext

    if stype == "pdf":
        doc = parse_pdf(content)
    elif stype == "docx":
        doc = parse_docx(content)
    elif stype in ("md", "markdown"):
        doc = parse_text(content, "md")
    else:
        doc = parse_text(content, "txt")

    doc.text = clean_text(doc.text)
    doc.language = detect_language(doc.text)
    return doc
